"""
Shared python-docx helpers used by all skill-generated Word document scripts.

Scripts are saved to Outputs/{TICKER}/ and run from the project root, so they
must add '.' to sys.path before importing:

    import sys; sys.path.insert(0, '.')
    from doc_utils import autofit_table, add_table_borders, set_row_font_size
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


def autofit_table(table):
    """Set table layout to autofit and strip all fixed w:tcW cell-width overrides."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for tag, attrs in [('w:tblW', {'w:w': '0', 'w:type': 'auto'}), ('w:tblLayout', {'w:type': 'autofit'})]:
        el = tblPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        if el not in list(tblPr):
            tblPr.append(el)
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                for tcW in tcPr.findall(qn('w:tcW')):
                    tcPr.remove(tcW)


def add_table_borders(table):
    """Apply a thin single border to all four sides (and inner dividers) of every cell."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)


def set_row_font_size(row, size=12):
    """Set font size (in points) for all runs in all cells of a table row."""
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(size)
